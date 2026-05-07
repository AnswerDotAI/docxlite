'''Read, render, search, and edit .docx files with tracked changes.

This skill is for LLM-friendly Word document editing inside notebooks and solveit dialogs.
It builds on `python-docx`, adding markdown previews, markdown-to-runs insertion,
document-order block iteration, raw-text search, tracked changes, and small editing
helpers for paragraphs, cells, rows, and tables.

Importing this skill enables tracked changes by default:

    set_tracking('AI Editor')

Use `set_tracking(None)` to disable revision tracking.

## Opening and rendering

    doc = DocxDocument('contract.docx')
    doc  # renders the full document as markdown

    p = doc.paragraphs[0]
    p  # renders one paragraph

    tbl = doc.tables[0]
    tbl  # renders one table as a markdown table

Markdown rendering is a readable preview, not a full docx serializer. It shows common
formatting and tracked insertions/deletions, but complex Word features may not appear
exactly as they do in Word.

## Document structure

A .docx body contains two top-level block types:

    Document -> Paragraph / Table
    Table -> Row -> Cell -> Paragraph -> Run

Paragraphs and tables are siblings in the document body. `doc.paragraphs` and
`doc.tables` lose that interleaving, so use document iteration when order matters.

## Document-order blocks

Documents are iterable. Iteration yields top-level paragraphs and tables in the order
they appear in the body.

    for block in doc:
        print(block)

    doc.blocks  # materialized list of top-level Paragraph/Table objects

The yielded objects are real `python-docx` objects patched by docxlite, so they can be
edited directly:

    block = doc.blocks[3]
    block.insert_after('New text')

## Search

`doc.search()` searches document-order blocks and returns the matching Paragraph/Table
objects themselves. It searches raw text (`._text`), not markdown preview text, so bold,
italic, underline, and revision spans do not affect matching.

Exact search is case-insensitive by default:

    hits = doc.search('termination')
    first = hits[0]
    first.insert_after('New language here.')

Regex search is also supported:

    hits = doc.search(r'terminat\\w+', regex=True)

Use `case=True` for case-sensitive matching:

    hits = doc.search('Effective Date', case=True)

Current search includes text from normal runs, inserted runs, and deleted runs. That is
usually what you want when reviewing tracked-change documents.

## Runs and revisions

Runs are the leaf text objects. docxlite patches `Paragraph.runs` so it includes runs
inside tracked insertions and deletions.

    r = p.runs[0]
    r.text
    r.revision  # Revision(typ='ins'/'del', author='...', date='...') or None

Tracked changes are represented with WordprocessingML elements:

    w:ins      inserted content
    w:del      deleted content
    w:delText  deleted run text

## Markdown insertion

`add_md` parses a small markdown subset and creates formatted runs.

Supported inline formatting:

    **bold**
    *italic*
    ***bold italic***
    <u>underline</u>
    <u>**bold underline**</u>

Examples:

    p.add_md('This is **bold** and *italic*')
    p.add_md('<u>underlined</u> text')

`Document.add_paragraph()` is patched to use markdown too:

    p = doc.add_paragraph('<u>**EXHIBIT C**</u>')

Unsupported markdown is not intended to round-trip perfectly. Keep inserted text simple.

## Paragraph base formatting

Use `Paragraph.apply_base(src)` to copy base paragraph formatting from another paragraph
onto an existing paragraph.

    exhibit = doc.search('EXHIBIT A', case=True)[0]
    p = doc.add_paragraph('<u>**EXHIBIT C**</u>')
    p.apply_base(exhibit)

`apply_base` copies paragraph-level formatting such as alignment, spacing, indentation,
tabs, and paragraph properties. It also applies base run defaults such as font family,
size, and color to existing runs.

It deliberately does not copy emphasis such as bold, italic, or underline. Use markdown
for those.

This makes call order forgiving:

    p = doc.add_paragraph('<u>**EXHIBIT C**</u>')
    p.apply_base(exhibit)

## Inserting paragraphs and tables

`insert_before` and `insert_after` work on both paragraphs and tables.

Passing a string inserts a paragraph:

    p2 = p.insert_after('New paragraph text')
    p0 = p.insert_before('Earlier paragraph')

When inserting a paragraph after/before another paragraph, docxlite applies the anchor
paragraph as the base automatically. This mirrors Word's local continuation behavior.

Passing `(rows, cols)` inserts a fresh table:

    tbl = p.insert_after((0, 3))          # empty 3-column table after paragraph
    tbl = p.insert_before((1, 2))         # 1-row, 2-column table before paragraph
    tbl = existing_tbl.insert_after((0, 4))

A table inserted after/before another block is a new table using Word/python-docx
defaults. It does not try to clone nearby table geometry.

When inserting a paragraph after a table, there is no paragraph context to continue.
Apply a base paragraph explicitly if needed:

    p = tbl.insert_after('<u>**EXHIBIT C**</u>')
    p.apply_base(exhibit)

## Tables

Tables render as markdown tables for preview. Merged cells are rendered without
duplicating the underlying python-docx layout-grid proxies.

Create a table using normal python-docx:

    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = 'Table Grid'

`Table.add_row` accepts markdown cell values:

    tbl.add_row('Name', 'Value', 'Notes')
    tbl.add_row(['Alpha', '42', '**First entry**'])

If fewer values are provided than there are columns, the final value may be merged across
the remaining cells, depending on the current docxlite patch behavior.

Cells support markdown insertion:

    tbl.cell(0, 0).add_md('**Name**')

Cells also support explicit merging:

    row = tbl.add_row()
    row.cells[0].merge(row.cells[2]).add_md('Merged heading')

For complex tables, prefer explicit structure:

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'

    r = tbl.add_row()
    r.cells[0].merge(r.cells[3]).add_md('**Centered header**')
    r.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    tbl.add_row('Col 1', 'Col 2', 'Col 3', 'Col 4')

    r = tbl.add_row('Vertical', 'Col 2', 'Col 3', 'Col 4')
    r.cells[0].merge(tbl.rows[-1].cells[0])

## Alignment and positioning

Use the normal python-docx enums for alignment. docxlite exports the common ones.

Paragraph horizontal alignment:

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

Cell vertical alignment:

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

Table block alignment between page margins:

    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT

Horizontal alignment inside a cell is paragraph alignment on the cell's paragraphs:

    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

## Deleting

Deletion methods are tracking-aware.

With tracking disabled, `.delete()` physically removes content.

With tracking enabled:

    run.delete()   # marks a normal run as deleted
    p.delete()     # marks all runs and the paragraph mark as deleted
    row.delete()   # marks the row and cell contents as deleted
    tbl.delete()   # marks all rows and contents as deleted

Deleting content that was itself a tracked insertion removes the insertion instead of
creating a deletion of an insertion.

## Formatting notes

Formatting inherits through Word's normal chain:

    Run -> Paragraph style -> Document defaults

Run-level properties include:

    run.font.bold
    run.font.italic
    run.font.underline
    run.font.name
    run.font.size
    run.font.strike

Use `Pt(12)` from `docx.shared` for font sizes.

Lists are paragraphs with numbering properties, not list containers. To add an item to
an existing list, insert before or after a nearby list paragraph so numbering properties
are copied.

## Saving

    doc.save('output.docx')
'''

from docx import Document as DocxDocument
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Emu, Inches, Pt
from docx.table import Table, _Row, _Cell
from docx.text.font import Font
from docx.text.paragraph import Paragraph
from docx.text.parfmt import ParagraphFormat
from docx.text.run import Run
from docxlite.core import *
from pyskills.core import allow
from safepyrun.core import allow_write_types

__all__ = [
    'DocxDocument', 'set_tracking',
    'Document', 'Paragraph', 'Run', 'Table', '_Row', '_Cell',
    'Revision', 'Underline',
    'Pt', 'Inches', 'Emu',
    'WD_ALIGN_PARAGRAPH', 'WD_CELL_VERTICAL_ALIGNMENT', 'WD_TABLE_ALIGNMENT',
]

set_tracking('Docxlite')
allow({_Cell: ..., _Row: ..., Document: ..., Paragraph: ..., Run: ..., Table: ...},
      DocxDocument, Emu, Inches, Pt, WD_ALIGN_PARAGRAPH, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT)
allow_write_types(_Cell, _Row, Font, Paragraph, ParagraphFormat, Table)
